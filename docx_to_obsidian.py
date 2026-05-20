#!/usr/bin/env python3
"""
docx_to_obsidian.py
====================
Leest een .docx of .pdf bestand en zet de inhoud om naar Obsidian-compatible
Markdown (.md) bestanden in een opgegeven vault-map.

Gebruik:
    python docx_to_obsidian.py boek.docx --vault ~/ObsidianVault/Books
    python docx_to_obsidian.py boek.docx --vault ~/ObsidianVault/Books --split-by heading1
    python docx_to_obsidian.py boek.docx --vault ~/ObsidianVault/Books --single-file
    python docx_to_obsidian.py input/boek.pdf --pages-per-section 5

Opties:
    --vault DIR              Map naar je Obsidian vault (standaard: ./RSDO)
    --split-by MODE          Alleen .docx: heading1 | heading2 | none (standaard: heading1)
    --pages-per-section N    Alleen .pdf: aantal pagina's per sectie (standaard: 5)
    --single-file            Alles in één .md bestand
    --prefix TAG             Obsidian tag toevoegen aan frontmatter (bijv. boek)
    --dry-run                Toon wat er gemaakt zou worden zonder te schrijven

Vereisten:
    pip install python-docx        # voor .docx
    pip install pdfplumber         # voor .pdf
"""

import argparse
import os
import re
import sys
from pathlib import Path
from datetime import date

# ── Controleer benodigde libraries ──────────────────────────────────────────
def check_deps(need_pdf: bool = False):
    missing = []
    if need_pdf:
        try:
            import pdfplumber  # noqa: F401
        except ImportError:
            missing.append("pdfplumber")
    else:
        try:
            import docx  # python-docx
        except ImportError:
            missing.append("python-docx")
    if missing:
        print(f"Ontbrekende libraries: {', '.join(missing)}")
        print(f"   Installeer met: pip install {' '.join(missing)}")
        sys.exit(1)

# ── Hulpfuncties ─────────────────────────────────────────────────────────────

def slugify(text: str) -> str:
    """Zet tekst om naar een veilige bestandsnaam."""
    text = text.strip().lower()
    text = re.sub(r"[^\w\s-]", "", text)
    text = re.sub(r"[\s_-]+", "-", text)
    text = re.sub(r"^-+|-+$", "", text)
    return text or "untitled"


def para_to_md(para) -> str:
    """Zet een docx Paragraph om naar Markdown-tekst."""
    style = para.style.name if para.style else ""
    text = para.text.strip()

    if not text:
        return ""

    # Koppen
    heading_map = {
        "Heading 1": "# ",
        "Heading 2": "## ",
        "Heading 3": "### ",
        "Heading 4": "#### ",
        "Heading 5": "##### ",
        "Heading 6": "###### ",
    }
    for h_style, md_prefix in heading_map.items():
        if style.startswith(h_style):
            return f"{md_prefix}{text}"

    # Lijsten
    if style.startswith("List Bullet"):
        level = int(style[-1]) if style[-1].isdigit() else 1
        indent = "  " * (level - 1)
        return f"{indent}- {text}"
    if style.startswith("List Number"):
        level = int(style[-1]) if style[-1].isdigit() else 1
        indent = "  " * (level - 1)
        return f"{indent}1. {text}"

    # Inline opmaak (bold/italic) per run
    md_runs = []
    for run in para.runs:
        run_text = run.text
        if not run_text:
            continue
        if run.bold and run.italic:
            run_text = f"***{run_text}***"
        elif run.bold:
            run_text = f"**{run_text}**"
        elif run.italic:
            run_text = f"*{run_text}*"
        md_runs.append(run_text)

    line = "".join(md_runs).strip()

    # Blockquote
    if style == "Quote" or style == "Intense Quote":
        return f"> {line}"

    return line


def docx_to_sections(doc_path: str, split_by: str):
    """
    Lees het .docx bestand en geef een lijst van (titel, [markdown_regels]) terug.
    split_by: 'heading1' | 'heading2' | 'none'
    """
    from docx import Document  # lazy: PDF-only gebruik vereist geen python-docx

    doc = Document(doc_path)
    book_title = Path(doc_path).stem

    if split_by == "none":
        lines = []
        for para in doc.paragraphs:
            md = para_to_md(para)
            lines.append(md)
        return [(book_title, lines)]

    split_heading = "Heading 1" if split_by == "heading1" else "Heading 2"
    sections = []
    current_title = book_title
    current_lines = []

    for para in doc.paragraphs:
        style = para.style.name if para.style else ""
        is_split = style.startswith(split_heading)

        if is_split:
            if current_lines or sections:
                sections.append((current_title, current_lines))
            current_title = para.text.strip() or "Zonder Titel"
            current_lines = []
        else:
            md = para_to_md(para)
            current_lines.append(md)

    # Laatste sectie
    if current_title or current_lines:
        sections.append((current_title, current_lines))

    return sections if sections else [(book_title, [])]


def pdf_to_sections(pdf_path: str, pages_per_section: int = 5):
    """
    Lees een .pdf bestand en geef een lijst van (titel, [tekstregels]) terug —
    dezelfde vorm als docx_to_sections, zodat de rest van het script ongewijzigd
    werkt (clean_lines, frontmatter, index, navigatie).

    PDF's hebben geen Heading-stijlen om op te splitsen, dus groeperen we pagina's:
    elke `pages_per_section` pagina's worden één sectie. Bij pages_per_section <= 0
    komt alles in één sectie.
    """
    import pdfplumber

    book_title = Path(pdf_path).stem
    sections = []
    with pdfplumber.open(pdf_path) as pdf:
        total = len(pdf.pages)
        step = pages_per_section if pages_per_section and pages_per_section > 0 else total or 1
        for start in range(0, total, step):
            end = min(start + step, total)
            title = book_title if step >= total else f"Pagina {start + 1}-{end}"
            lines = []
            for page in pdf.pages[start:end]:
                text = page.extract_text() or ""
                lines.extend(text.splitlines())
                lines.append("")  # lege regel als pagina-scheiding
            sections.append((title, lines))
    return sections if sections else [(book_title, [])]


def build_frontmatter(title: str, book_title: str, tags: list, index: int, total: int) -> str:
    """Maak Obsidian YAML frontmatter."""
    tag_str = "\n".join(f"  - {t}" for t in tags)
    return f"""---
title: "{title}"
book: "{book_title}"
tags:
{tag_str}
created: {date.today().isoformat()}
chapter: {index}
total_chapters: {total}
---

"""


def clean_lines(lines: list) -> str:
    """Verwijder overtollige lege regels en geef nette tekst terug."""
    result = []
    prev_empty = False
    for line in lines:
        if line == "":
            if not prev_empty:
                result.append("")
            prev_empty = True
        else:
            result.append(line)
            prev_empty = False
    return "\n".join(result).strip()


# ── Hoofdprogramma ────────────────────────────────────────────────────────────

def main():
    # Windows-console (cp1252) kan emoji niet printen → forceer UTF-8 uitvoer.
    try:
        sys.stdout.reconfigure(encoding="utf-8")
    except Exception:
        pass

    parser = argparse.ArgumentParser(
        description="Zet een .docx of .pdf boek om naar Obsidian Markdown notities."
    )
    parser.add_argument("input_file", help="Pad naar het .docx of .pdf bestand")
    default_vault = Path(__file__).parent / "RSDO"
    parser.add_argument(
        "--vault",
        default=str(default_vault),
        help=f"Pad naar je Obsidian vault map (standaard: {default_vault})"
    )
    parser.add_argument(
        "--split-by",
        choices=["heading1", "heading2", "none"],
        default="heading1",
        help="Splits het boek per Heading 1, Heading 2, of niet splitsen (standaard: heading1)"
    )
    parser.add_argument(
        "--single-file",
        action="store_true",
        help="Zet alles in één .md bestand"
    )
    parser.add_argument(
        "--prefix",
        default="boek",
        help="Obsidian tag voor het frontmatter (standaard: boek)"
    )
    parser.add_argument(
        "--pages-per-section",
        type=int,
        default=5,
        help="Alleen voor PDF: aantal pagina's per sectie (standaard: 5; <=0 = alles in één)"
    )
    parser.add_argument(
        "--dry-run",
        action="store_true",
        help="Toon wat er gemaakt zou worden zonder te schrijven"
    )
    args = parser.parse_args()

    # Validatie
    input_path = Path(args.input_file).expanduser().resolve()
    if not input_path.exists():
        print(f"Bestand niet gevonden: {input_path}")
        sys.exit(1)

    suffix = input_path.suffix.lower()
    if suffix not in (".docx", ".pdf"):
        print(f"Niet-ondersteund bestandstype: {suffix} (gebruik .docx of .pdf)")
        sys.exit(1)

    check_deps(need_pdf=(suffix == ".pdf"))

    vault_path = Path(args.vault).expanduser().resolve()
    book_title = input_path.stem
    output_dir = vault_path / slugify(book_title)

    if suffix == ".pdf":
        split_mode = "pages"
    else:
        split_mode = "none" if args.single_file else args.split_by

    print(f"\nBoek: {book_title}")
    print(f"Vault map: {output_dir}")
    print(f"Splitsing: {split_mode}")
    print(f"Tag: {args.prefix}\n")

    # Lees het bronbestand
    print(f"Bezig met lezen van het {suffix} bestand...")
    try:
        if suffix == ".pdf":
            sections = pdf_to_sections(str(input_path), args.pages_per_section)
        else:
            sections = docx_to_sections(str(input_path), split_mode)
    except Exception as e:
        print(f"Fout bij het lezen: {e}")
        sys.exit(1)

    total = len(sections)
    print(f"✅ {total} sectie(s) gevonden\n")

    # Maak de output map
    if not args.dry_run:
        output_dir.mkdir(parents=True, exist_ok=True)

    # Schrijf de bestanden
    tags = [args.prefix, slugify(book_title)]
    created_files = []

    if args.single_file:
        filename = f"{slugify(book_title)}.md"
        filepath = output_dir / filename
        all_content = []
        fm = build_frontmatter(book_title, book_title, tags, 1, 1)
        all_content.append(fm)
        all_content.append(f"# {book_title}\n")
        for title, lines in sections:
            if title != book_title:
                all_content.append(f"\n## {title}\n")
            all_content.append(clean_lines(lines))
        content = "\n\n".join(all_content)

        if args.dry_run:
            print(f"  [DRY RUN] → {filepath}")
            print(content[:500] + "...\n")
        else:
            filepath.write_text(content, encoding="utf-8")
            print(f"  ✅ {filepath}")
        created_files.append(filepath)
    else:
        # Index bestand
        index_path = output_dir / f"{slugify(book_title)}-index.md"
        index_lines = [
            f"---\ntitle: \"{book_title} - Index\"\ntags:\n  - {args.prefix}\n  - index\ncreated: {date.today().isoformat()}\n---\n",
            f"# {book_title}\n",
            "## Hoofdstukken\n",
        ]

        for i, (title, lines) in enumerate(sections, 1):
            filename = f"{i:02d}-{slugify(title)}.md"
            filepath = output_dir / filename
            fm = build_frontmatter(title, book_title, tags, i, total)
            # Wikilink naar vorige/volgende
            nav = []
            if i > 1:
                prev_title = sections[i - 2][0]
                prev_file = f"{(i-1):02d}-{slugify(prev_title)}"
                nav.append(f"← [[{prev_file}|{prev_title}]]")
            if i < total:
                next_title = sections[i][0]
                next_file = f"{(i+1):02d}-{slugify(next_title)}"
                nav.append(f"[[{next_file}|{next_title}]] →")
            nav_str = "   ".join(nav)

            content_parts = [fm]
            if nav_str:
                content_parts.append(nav_str + "\n\n---\n")
            content_parts.append(f"# {title}\n")
            body = clean_lines(lines)
            if body:
                content_parts.append(body)
            if nav_str:
                content_parts.append(f"\n---\n{nav_str}")
            content = "\n".join(content_parts)

            if args.dry_run:
                print(f"  [DRY RUN] → {filepath}")
            else:
                filepath.write_text(content, encoding="utf-8")
                print(f"  ✅ {filename}")

            created_files.append(filepath)
            # Voeg toe aan index
            index_lines.append(f"{i}. [[{filename[:-3]}|{title}]]")

        # Schrijf index
        index_content = "\n".join(index_lines)
        if args.dry_run:
            print(f"\n  [DRY RUN] → {index_path}")
        else:
            index_path.write_text(index_content, encoding="utf-8")
            print(f"\n  ✅ Index: {index_path.name}")
        created_files.append(index_path)

    print(f"\n🎉 Klaar! {len(created_files)} bestand(en) aangemaakt in:")
    print(f"   {output_dir}\n")
    print("💡 Tips voor Obsidian:")
    print("   • Open de vault map in Obsidian als vault")
    print("   • Gebruik Graph View om verbindingen tussen hoofdstukken te zien")
    print(f"   • Filter op tag #{args.prefix} om alle boeknotities te vinden")


if __name__ == "__main__":
    main()
